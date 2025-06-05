import asyncio
import os
import re
from typing import Optional, List
import docx
import PyPDF2
import pysrt
from sqlalchemy.orm import Session
import structlog

from app.core.database import SessionLocal
from app.models.file import File, FileChunk
from app.services.translation_service import translation_service

logger = structlog.get_logger()

class FileService:
    """文件处理服务"""
    
    def __init__(self):
        self.chunk_size = 1000  # 每个分块的字符数
    
    async def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """从文件中提取文本"""
        try:
            if file_type == "txt":
                return await self._extract_from_txt(file_path)
            elif file_type == "docx":
                return await self._extract_from_docx(file_path)
            elif file_type == "pdf":
                return await self._extract_from_pdf(file_path)
            elif file_type == "srt":
                return await self._extract_from_srt(file_path)
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
        
        except Exception as e:
            logger.error("Failed to extract text from file", file_path=file_path, file_type=file_type, error=str(e))
            raise
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """从TXT文件提取文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            encodings = ['gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法识别文件编码")
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """从DOCX文件提取文本"""
        doc = docx.Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        return '\n'.join(text_parts)
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        text_parts = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
        
        return '\n'.join(text_parts)
    
    async def _extract_from_srt(self, file_path: str) -> str:
        """从SRT字幕文件提取文本"""
        subs = pysrt.open(file_path, encoding='utf-8')
        text_parts = []
        
        for sub in subs:
            if sub.text.strip():
                # 清理HTML标签
                clean_text = re.sub(r'<[^>]+>', '', sub.text)
                text_parts.append(clean_text)
        
        return '\n'.join(text_parts)
    
    def split_text_into_chunks(self, text: str, chunk_size: int = None) -> List[str]:
        """将文本分割成块"""
        if chunk_size is None:
            chunk_size = self.chunk_size
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        sentences = re.split(r'[.!?。！？]\s*', text)
        current_chunk = ""
        
        for sentence in sentences:
            if not sentence.strip():
                continue
            
            # 如果添加这个句子会超过块大小，保存当前块并开始新块
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += ". " + sentence
                else:
                    current_chunk = sentence
        
        # 添加最后一个块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    async def translate_file_async(self, file_id: int, db: Session = None):
        """异步翻译文件"""
        if db is None:
            db = SessionLocal()
        
        try:
            # 获取文件记录
            file_record = db.query(File).filter(File.id == file_id).first()
            if not file_record:
                logger.error("File not found", file_id=file_id)
                return
            
            logger.info("Starting file translation", file_id=file_id, filename=file_record.original_filename)
            
            # 提取文本
            if not file_record.extracted_text:
                try:
                    extracted_text = await self.extract_text_from_file(
                        file_record.file_path, 
                        file_record.file_type
                    )
                    file_record.extracted_text = extracted_text
                    db.commit()
                except Exception as e:
                    file_record.translation_status = "failed"
                    file_record.error_message = f"文本提取失败: {str(e)}"
                    db.commit()
                    return
            
            # 分割文本
            chunks = self.split_text_into_chunks(file_record.extracted_text)
            total_chunks = len(chunks)
            
            # 创建分块记录
            for i, chunk_text in enumerate(chunks):
                chunk = FileChunk(
                    file_id=file_id,
                    chunk_index=i,
                    chunk_text=chunk_text,
                    status="pending"
                )
                db.add(chunk)
            
            db.commit()
            
            # 翻译每个分块
            translated_chunks = []
            for i, chunk_text in enumerate(chunks):
                try:
                    # 更新进度
                    progress = int((i / total_chunks) * 100)
                    file_record.translation_progress = progress
                    db.commit()
                    
                    # 翻译分块
                    result = await translation_service.translate(
                        text=chunk_text,
                        source_lang=file_record.source_language,
                        target_lang=file_record.target_language,
                        engine=file_record.translation_engine
                    )
                    
                    translated_text = result["translated_text"]
                    translated_chunks.append(translated_text)
                    
                    # 更新分块状态
                    chunk_record = db.query(FileChunk).filter(
                        FileChunk.file_id == file_id,
                        FileChunk.chunk_index == i
                    ).first()
                    
                    if chunk_record:
                        chunk_record.translated_text = translated_text
                        chunk_record.status = "completed"
                        db.commit()
                    
                    logger.info(f"Chunk {i+1}/{total_chunks} translated", file_id=file_id)
                    
                except Exception as e:
                    logger.error(f"Failed to translate chunk {i}", file_id=file_id, error=str(e))
                    
                    # 更新分块状态为失败
                    chunk_record = db.query(FileChunk).filter(
                        FileChunk.file_id == file_id,
                        FileChunk.chunk_index == i
                    ).first()
                    
                    if chunk_record:
                        chunk_record.status = "failed"
                        db.commit()
                    
                    # 使用原文作为翻译结果
                    translated_chunks.append(chunk_text)
            
            # 合并翻译结果
            full_translated_text = '\n'.join(translated_chunks)
            file_record.translated_content = full_translated_text
            
            # 生成翻译后的文件
            translated_file_path = await self._generate_translated_file(
                file_record, full_translated_text
            )
            
            file_record.translated_file_path = translated_file_path
            file_record.translation_status = "completed"
            file_record.translation_progress = 100
            
            from datetime import datetime
            file_record.completed_at = datetime.utcnow()
            
            db.commit()
            
            logger.info("File translation completed", file_id=file_id)
            
        except Exception as e:
            logger.error("File translation failed", file_id=file_id, error=str(e))
            
            file_record.translation_status = "failed"
            file_record.error_message = str(e)
            db.commit()
        
        finally:
            db.close()
    
    async def _generate_translated_file(self, file_record: File, translated_text: str) -> str:
        """生成翻译后的文件"""
        try:
            # 生成输出文件路径
            base_name = os.path.splitext(file_record.filename)[0]
            output_filename = f"{base_name}_translated.{file_record.file_type}"
            output_path = os.path.join(
                os.path.dirname(file_record.file_path),
                output_filename
            )
            
            if file_record.file_type == "txt":
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(translated_text)
            
            elif file_record.file_type == "docx":
                # 创建新的Word文档
                doc = docx.Document()
                for paragraph_text in translated_text.split('\n'):
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text)
                doc.save(output_path)
            
            elif file_record.file_type == "srt":
                # 重新构建SRT文件
                await self._generate_translated_srt(
                    file_record.file_path,
                    output_path,
                    translated_text
                )
            
            else:
                # 对于PDF等其他格式，生成TXT文件
                output_path = output_path.replace(f'.{file_record.file_type}', '.txt')
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(translated_text)
            
            return output_path
            
        except Exception as e:
            logger.error("Failed to generate translated file", error=str(e))
            raise
    
    async def _generate_translated_srt(self, original_path: str, output_path: str, translated_text: str):
        """生成翻译后的SRT字幕文件"""
        try:
            # 读取原始SRT文件
            original_subs = pysrt.open(original_path, encoding='utf-8')
            
            # 分割翻译文本
            translated_lines = [line.strip() for line in translated_text.split('\n') if line.strip()]
            
            # 创建新的字幕文件
            new_subs = pysrt.SubRipFile()
            
            for i, sub in enumerate(original_subs):
                if i < len(translated_lines):
                    new_sub = pysrt.SubRipItem(
                        index=sub.index,
                        start=sub.start,
                        end=sub.end,
                        text=translated_lines[i]
                    )
                    new_subs.append(new_sub)
                else:
                    # 如果翻译文本不够，使用原文
                    new_subs.append(sub)
            
            # 保存新的SRT文件
            new_subs.save(output_path, encoding='utf-8')
            
        except Exception as e:
            logger.error("Failed to generate translated SRT", error=str(e))
            # 如果SRT生成失败，创建简单的文本文件
            with open(output_path.replace('.srt', '.txt'), 'w', encoding='utf-8') as f:
                f.write(translated_text)

# 创建全局文件服务实例
file_service = FileService()